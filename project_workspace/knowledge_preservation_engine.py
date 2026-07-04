import json
import hashlib
from pathlib import Path

from docx import Document

from project_workspace.pdf_acquisition_engine import (
    PDFAcquisitionEngine
)


class KnowledgePreservationEngine:
    """
    AROS Knowledge Preservation Engine v3.

    Domain neutral engine.

    Works for:
    - Finance
    - Accounting
    - Governance
    - Taxation
    - Economics
    - Management
    - AI Research

    Every research object becomes:

    1. PDF (if available)
    2. JSON knowledge record
    3. DOCX research card
    """

    def __init__(self):

        self.pdf_engine = PDFAcquisitionEngine()


    def preserve(
        self,
        paper,
        output_type,
        project_id,
        domain,
        impact_factor="NA"
    ):

        acquisition = self.pdf_engine.acquire(
            paper=paper,
            output_type=output_type,
            project_id=project_id,
            domain=domain,
            impact_factor=impact_factor
        )


        card = self.create_research_card(
            paper,
            output_type,
            project_id
        )


        return {
            "knowledge_status": "Preserved",
            "pdf_or_metadata": acquisition,
            "research_card": card
        }



    def create_research_card(
        self,
        paper,
        output_type,
        project_id
    ):

        folder = (
            Path("Research_Output")
            / output_type
            / "Researched_Library"
            / project_id
            / "03_Research_Cards"
        )


        folder.mkdir(
            parents=True,
            exist_ok=True
        )


        unique_text = (
            (paper.doi or "")
            + (paper.title or "")
            + (paper.source or "")
            + (paper.publication_year or "")
        )

        unique_id = hashlib.md5(
            unique_text.encode("utf-8")
        ).hexdigest()[:8]

        clean_title = (
            (paper.title or "untitled")[:60]
            .replace("/", "_")
            .replace(" ", "_")
        )

        filename = (
            f"{paper.publication_year or 'Year_NA'}_"
            f"{paper.source or 'Source_NA'}_"
            f"{clean_title}_"
            f"{unique_id}.docx"
        )


        path = folder / filename


        doc = Document()


        doc.add_heading(
            paper.title,
            level=1
        )


        fields = {

            "Source":
                paper.source,

            "Year":
                paper.publication_year,

            "DOI":
                paper.doi,

            "Authors":
                ", ".join(
                    paper.authors
                ),

            "Research Domain":
                paper.research_domain,

            "Abstract":
                paper.abstract,


            "AROS Research Notes":
                (
                    "Theory:\n\n"
                    "Variables:\n\n"
                    "Methodology:\n\n"
                    "Findings:\n\n"
                    "Limitations:\n\n"
                    "Possible Research Gap:\n"
                )

        }


        for key, value in fields.items():

            doc.add_heading(
                key,
                level=2
            )

            doc.add_paragraph(
                str(value or "")
            )


        doc.save(path)


        return str(path)

